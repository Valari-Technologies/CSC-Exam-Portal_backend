"""File parsers for question bank bulk import.

Supported formats: .csv, .xlsx, .xls, .json, .txt, .docx
Expected columns: question_text, option_a, option_b, option_c, option_d,
    correct_option, subject_id, chapter_id, difficulty, marks
Optional columns: explanation, negative_marks

`subject_id` carries the generated Subject ID (e.g. KA_MAT_10), not a database key.
It replaced `subject_name` on 2026-07-20: subject names are only unique per class, so
"Mathematics" matches ten rows in a school and name lookups were nearly always ambiguous.
"""
import csv
import io
import json
import logging
from typing import Dict, List

logger = logging.getLogger(__name__)

REQUIRED_COLUMNS = {
    'question_text',
    'option_a',
    'option_b',
    'option_c',
    'option_d',
    'correct_option',
    'subject_id',
    'chapter_id',
    'difficulty',
    'marks',
}

# Columns that must always be present. `subject_id` is required outright — the Subject ID
# identifies a subject exactly, so unlike chapters there is no name-based alternative.
REQUIRED_COLUMNS_FIXED = {
    'question_text',
    'option_a', 'option_b', 'option_c', 'option_d',
    'correct_option',
    'subject_id',
    'difficulty',
    'marks',
}

# A chapter may still be supplied by id OR by name — chapter names are unique within their
# subject, so that lookup is unambiguous once the subject is resolved.
CHAPTER_COLUMNS = {'chapter_id', 'chapter_name'}

OPTIONAL_COLUMNS = {'explanation', 'negative_marks'}


def validate_required_columns(present_columns: set) -> set:
    """Return the set of missing required columns.

    For chapter, at least ONE of the id/name pair must be present.
    """
    missing = REQUIRED_COLUMNS_FIXED - present_columns
    if not (CHAPTER_COLUMNS & present_columns):
        missing.add('chapter_id or chapter_name')
    return missing


SUPPORTED_EXTENSIONS = ('.csv', '.xlsx', '.xls', '.json', '.txt', '.docx')


def parse_questions_file(file_obj) -> List[Dict[str, str]]:
    """Parse an uploaded file and return a list of row dicts.

    Args:
        file_obj: Django ``UploadedFile`` (or any file-like with ``.name``).

    Returns:
        List of dicts where keys are column headers.

    Raises:
        ValueError: if the file format is unsupported or cannot be parsed.
    """
    name = file_obj.name.lower()

    if name.endswith('.csv'):
        return _parse_csv(file_obj)
    elif name.endswith('.xlsx') or name.endswith('.xls'):
        return _parse_excel(file_obj)
    elif name.endswith('.json'):
        return _parse_json(file_obj)
    elif name.endswith('.txt'):
        return _parse_txt(file_obj)
    elif name.endswith('.docx'):
        return _parse_docx(file_obj)
    else:
        raise ValueError(f'Unsupported file format: {name}')


def _parse_csv(file_obj) -> List[Dict[str, str]]:
    """Parse a CSV file using csv.DictReader."""
    file_obj.seek(0)
    decoded = file_obj.read().decode('utf-8-sig')
    reader = csv.DictReader(io.StringIO(decoded))
    return list(reader)


def _parse_excel(file_obj) -> List[Dict[str, str]]:
    """Parse .xlsx / .xls using openpyxl. First row treated as headers."""
    import openpyxl

    file_obj.seek(0)
    wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
    ws = wb.active
    if ws is None:
        raise ValueError('The Excel file has no active worksheet.')

    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    if not rows:
        return []

    headers = [str(h).strip() if h is not None else '' for h in rows[0]]
    result: List[Dict[str, str]] = []
    for row in rows[1:]:
        row_dict: Dict[str, str] = {}
        for idx, cell_value in enumerate(row):
            if idx < len(headers) and headers[idx]:
                row_dict[headers[idx]] = str(cell_value).strip() if cell_value is not None else ''
        if any(v for v in row_dict.values()):
            result.append(row_dict)

    return result


def _parse_json(file_obj) -> List[Dict[str, str]]:
    """Parse a JSON file containing an array of question objects."""
    file_obj.seek(0)
    decoded = file_obj.read().decode('utf-8-sig')
    try:
        data = json.loads(decoded)
    except json.JSONDecodeError as exc:
        raise ValueError(f'Invalid JSON file: {exc}')

    if not isinstance(data, list):
        raise ValueError('JSON file must contain an array of question objects.')

    result: List[Dict[str, str]] = []
    for item in data:
        if not isinstance(item, dict):
            raise ValueError('Each item in the JSON array must be an object.')
        row_dict: Dict[str, str] = {
            str(k).strip(): str(v).strip() if v is not None else ''
            for k, v in item.items()
        }
        result.append(row_dict)

    return result


def _parse_txt(file_obj) -> List[Dict[str, str]]:
    """Parse a tab-delimited or comma-delimited text file. First line is headers."""
    file_obj.seek(0)
    decoded = file_obj.read().decode('utf-8-sig')
    lines = decoded.strip().splitlines()

    if not lines:
        return []

    # Auto-detect delimiter: tab vs comma
    header_line = lines[0]
    delimiter = '\t' if '\t' in header_line else ','

    reader = csv.DictReader(io.StringIO(decoded), delimiter=delimiter)
    return list(reader)


def _parse_docx(file_obj) -> List[Dict[str, str]]:
    """Parse a .docx file by extracting tables. First row of each table is headers."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError(
            'python-docx is required to parse .docx files. '
            'Install it with: pip install python-docx'
        )

    file_obj.seek(0)
    doc = Document(file_obj)

    if not doc.tables:
        raise ValueError('No tables found in the Word document. Please structure questions in a table.')

    result: List[Dict[str, str]] = []
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        headers = [cell.text.strip() for cell in rows[0].cells]
        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            row_dict: Dict[str, str] = {}
            for idx, cell_value in enumerate(cells):
                if idx < len(headers) and headers[idx]:
                    row_dict[headers[idx]] = cell_value
            if any(v for v in row_dict.values()):
                result.append(row_dict)

    return result
