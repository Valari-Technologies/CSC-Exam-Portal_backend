"""File parsers for multi-format student bulk import.

Supported formats: .csv, .xlsx, .xls, .txt, .doc, .docx

Every parser returns rows keyed by *normalized* headers (see ``normalize_header``), so the
importer can match on ``full_name`` whether the admin typed "Full Name", "full name", or
"FULL_NAME ". Without this, a file built by hand in Excel is rejected with "Missing
columns: full_name" while the column is plainly there — the header just wasn't spelled the
way the code expected.
"""
import csv
import io
import logging
import re
from typing import Dict, List

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = ('.csv', '.xlsx', '.xls', '.txt', '.doc', '.docx')

_NON_ALNUM_RE = re.compile(r'[^a-z0-9]+')


def normalize_header(header: object) -> str:
    """Fold a spreadsheet column heading to the canonical snake_case key.

    'Full Name' / 'full name' / ' FULL_NAME ' / 'Full-Name' all become 'full_name'.
    """
    text = str(header if header is not None else '').strip().lower()
    return _NON_ALNUM_RE.sub('_', text).strip('_')


def _clean_cell(value: object) -> str:
    """Cell value as a trimmed string. Excel hands back ints/floats/None, not text."""
    if value is None:
        return ''
    if isinstance(value, float) and value.is_integer():
        # Excel stores a roll number typed as `12` as 12.0; '12.0' is not the roll number.
        value = int(value)
    return str(value).strip()


def _build_rows(headers: List[str], raw_rows) -> List[Dict[str, str]]:
    """Zip normalized headers against data rows, dropping rows that are entirely blank.

    Blank rows matter: saving a sheet as CSV from Excel routinely leaves a trailing empty
    line, and counting it as a failed student made a clean import report "1 failed".
    """
    result: List[Dict[str, str]] = []
    for row in raw_rows:
        row_dict = {
            headers[idx]: _clean_cell(cell)
            for idx, cell in enumerate(row)
            if idx < len(headers) and headers[idx]
        }
        if any(row_dict.values()):
            result.append(row_dict)
    return result


def parse_upload_file(file_obj) -> List[Dict[str, str]]:
    """Parse an uploaded file and return a list of row dicts.

    Args:
        file_obj: Django ``UploadedFile`` (or any file-like with ``.name``).

    Returns:
        List of dicts where keys are column headers (from the first row).

    Raises:
        ValueError: if the file format is unsupported or cannot be parsed.
    """
    name = file_obj.name.lower()

    if name.endswith('.csv'):
        return _parse_csv(file_obj)
    elif name.endswith('.xlsx'):
        return _parse_excel(file_obj)
    elif name.endswith('.txt'):
        return _parse_txt(file_obj)
    elif name.endswith('.docx'):
        return _parse_docx(file_obj)
    elif name.endswith('.xls'):
        # openpyxl reads .xlsx only — a real .xls is an OLE2 file and blows up inside the
        # zip reader. This was advertised as supported and never worked; say so plainly
        # rather than crashing.
        raise ValueError(
            'Legacy .xls format is not supported. Open the file in Excel and use '
            'File > Save As to save it as .xlsx (or CSV), then upload it again.'
        )
    elif name.endswith('.doc'):
        raise ValueError(
            'Legacy .doc format is not supported. '
            'Please save the file as .docx and re-upload.'
        )
    else:
        raise ValueError(f'Unsupported file format: {name}')


def _parse_delimited(decoded: str, dialect) -> List[Dict[str, str]]:
    """Shared CSV/TXT body: first row is headers, the rest are data."""
    rows = list(csv.reader(io.StringIO(decoded), dialect))
    if not rows:
        return []
    headers = [normalize_header(h) for h in rows[0]]
    return _build_rows(headers, rows[1:])


def _parse_csv(file_obj) -> List[Dict[str, str]]:
    """Parse a CSV file."""
    file_obj.seek(0)
    decoded = file_obj.read().decode('utf-8-sig')
    return _parse_delimited(decoded, csv.excel)


def _parse_excel(file_obj) -> List[Dict[str, str]]:
    """Parse .xlsx / .xls using openpyxl. First row is treated as headers."""
    import openpyxl

    file_obj.seek(0)
    wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
    ws = wb.active
    if ws is None:
        raise ValueError('The Excel file has no active worksheet.')

    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    if not rows:
        return []

    headers = [normalize_header(h) for h in rows[0]]
    return _build_rows(headers, rows[1:])


def _parse_txt(file_obj) -> List[Dict[str, str]]:
    """Parse a .txt file assuming tab-delimited or comma-delimited content.

    Uses csv.Sniffer to auto-detect the delimiter.
    """
    file_obj.seek(0)
    decoded = file_obj.read().decode('utf-8-sig')

    # Try to sniff the dialect from a sample
    sample = decoded[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters='\t,')
    except csv.Error:
        # Fall back to tab, then comma
        if '\t' in sample:
            dialect = csv.excel_tab
        else:
            dialect = csv.excel

    return _parse_delimited(decoded, dialect)


def _parse_docx(file_obj) -> List[Dict[str, str]]:
    """Parse a .docx file by extracting the first table found.

    First row is treated as headers, remaining rows as data.
    """
    import docx

    file_obj.seek(0)
    document = docx.Document(file_obj)

    if not document.tables:
        raise ValueError('No tables found in the Word document.')

    table = document.tables[0]
    rows = table.rows
    if len(rows) < 2:
        raise ValueError('The table must have at least a header row and one data row.')

    headers = [normalize_header(cell.text) for cell in rows[0].cells]
    return _build_rows(headers, ([cell.text for cell in row.cells] for row in rows[1:]))
